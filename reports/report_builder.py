"""
reports/report_builder.py
Two page report
Sections:
  1. Cover page: title, date, analyst, classification
  2. Executive Summary: per-query narrative + overall risk rating
  3. Event Data Tables: full results per query
  4. Flagged Records: auto-detected high-severity / suspicious events
  5. Attack Timeline: chart (embedded as image)
  6. Recommended Actions: LLM-generated per event type
  7. Appendix: DSL queries used
Export formats: .pdf- (reportlab) and .docx (python-docx)
"""

import io
import json
import os
import sys
from collections import Counter
from datetime import datetime

import pandas as pd

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))


# ── Recommendations ────────────────────────────────────────────────────────────

_RECS = {
    "brute_force": [
        ("IMMEDIATE",  "Block source IPs conducting brute force at the perimeter firewall"),
        ("IMMEDIATE",  "Lock accounts that exceeded the failed login threshold"),
        ("SHORT-TERM", "Implement account lockout: 5 attempts / 15 min lockout"),
        ("SHORT-TERM", "Enable geo-blocking for countries with no business presence"),
        ("LONG-TERM",  "Deploy rate limiting and CAPTCHA on all authentication endpoints"),
        ("LONG-TERM",  "Enforce MFA organisation-wide — prioritise privileged accounts"),
    ],
    "failed_login": [
        ("IMMEDIATE",  "Investigate accounts with high failure counts — possible credential stuffing"),
        ("SHORT-TERM", "Enforce MFA on all user accounts"),
        ("SHORT-TERM", "Audit and rotate credentials for affected accounts"),
        ("LONG-TERM",  "Implement password complexity and rotation policy"),
    ],
    "malware_detection": [
        ("IMMEDIATE",  "Isolate affected hosts from the network immediately"),
        ("IMMEDIATE",  "Preserve forensic images before any remediation"),
        ("SHORT-TERM", "Run full AV/EDR scan on all affected systems"),
        ("LONG-TERM",  "Deploy application whitelisting on critical systems"),
    ],
    "vpn_login": [
        ("IMMEDIATE",  "Review VPN logins from high-risk countries"),
        ("SHORT-TERM", "Enforce MFA on all VPN authentication"),
        ("SHORT-TERM", "Alert on concurrent sessions from different geolocations"),
        ("LONG-TERM",  "Implement zero-trust network access (ZTNA)"),
    ],
    "privilege_escalation": [
        ("IMMEDIATE",  "Revoke suspicious elevated sessions and review audit logs"),
        ("SHORT-TERM", "Audit and remove unnecessary sudo/admin privileges"),
        ("LONG-TERM",  "Implement just-in-time (JIT) access for all admin operations"),
    ],
    "data_exfiltration": [
        ("IMMEDIATE",  "Block outbound transfers to flagged destination IPs/countries"),
        ("IMMEDIATE",  "Preserve network logs for forensic investigation"),
        ("SHORT-TERM", "Implement Data Loss Prevention (DLP) policies"),
        ("LONG-TERM",  "Deploy UEBA to detect anomalous data access patterns"),
    ],
    "lateral_movement": [
        ("IMMEDIATE",  "Isolate affected network segments"),
        ("SHORT-TERM", "Disable unnecessary SMB, WMI, RDP between endpoints"),
        ("LONG-TERM",  "Implement microsegmentation across the environment"),
    ],
    "port_scan": [
        ("IMMEDIATE",  "Block source IPs conducting reconnaissance at the firewall"),
        ("SHORT-TERM", "Harden and reduce attack surface of exposed services"),
        ("LONG-TERM",  "Enable IDS/IPS rules for network scanning signatures"),
    ],
    "suspicious_powershell": [
        ("IMMEDIATE",  "Investigate the host running suspicious PowerShell"),
        ("SHORT-TERM", "Enable PowerShell Script Block Logging"),
        ("LONG-TERM",  "Deploy EDR with PowerShell-specific detection rules"),
    ],
    "file_integrity": [
        ("IMMEDIATE",  "Investigate unexpected changes to critical system files"),
        ("SHORT-TERM", "Review and tighten file system access controls"),
        ("LONG-TERM",  "Expand FIM coverage to all critical system directories"),
    ],
    "default": [
        ("SHORT-TERM", "Review affected systems and apply available patches"),
        ("SHORT-TERM", "Ensure security monitoring coverage is complete"),
        ("LONG-TERM",  "Conduct a full security review of affected assets"),
    ],
}


# ── Helpers ────────────────────────────────────────────────────────────────────

def _severity_label(level) -> str:
    try:
        lvl = int(level)
        if lvl >= 14: return "CRITICAL"
        if lvl >= 10: return "HIGH"
        if lvl >= 5:  return "MEDIUM"
        return "LOW"
    except Exception:
        return str(level)


def _risk_rating(queries_data: list[dict]) -> str:
    max_level = 0
    for q in queries_data:
        for hit in q.get("hits", []):
            try:
                lvl = int(hit.get("rule", {}).get("level", 0) or 0)
                max_level = max(max_level, lvl)
            except Exception:
                pass
    if max_level >= 14: return "CRITICAL"
    if max_level >= 10: return "HIGH"
    if max_level >= 5:  return "MEDIUM"
    return "LOW"


def _flag_records(hits: list[dict]) -> list[dict]:
    flagged    = []
    src_counts = Counter(
        h.get("data", {}).get("srcip", "") for h in hits
        if h.get("data", {}).get("srcip")
    )
    for h in hits:
        reasons = []
        level   = int(h.get("rule", {}).get("level", 0) or 0)
        if level >= 14:   reasons.append("Critical severity")
        elif level >= 12: reasons.append("High severity")
        src_ip = h.get("data", {}).get("srcip", "")
        if src_ip and src_counts[src_ip] >= 3:
            reasons.append(f"Repeated source IP ({src_counts[src_ip]}x)")
        desc = h.get("rule", {}).get("description", "").lower()
        for pattern, label in [
            ("mimikatz",         "Credential theft tool"),
            ("pass-the-hash",    "Pass-the-Hash"),
            ("lateral movement", "Lateral movement"),
            ("exfiltration",     "Data exfiltration"),
            ("ransomware",       "Ransomware"),
            ("privilege",        "Privilege escalation"),
            ("powershell",       "Suspicious PowerShell"),
        ]:
            if pattern in desc:
                reasons.append(label)
        if reasons:
            flagged.append({**h, "_flag_reasons": list(dict.fromkeys(reasons))})
    return flagged


def _compute_stats(queries_data: list[dict]) -> dict:
    all_hits  = [h for q in queries_data for h in q.get("hits", [])]
    sev_dist  = Counter(_severity_label(h.get("rule", {}).get("level", "")) for h in all_hits)
    et_dist   = Counter(h.get("event_type", "unknown") for h in all_hits)
    top_ips   = Counter(h.get("data", {}).get("srcip","") for h in all_hits
                        if h.get("data", {}).get("srcip")).most_common(5)
    top_users = Counter(h.get("user", {}).get("name","") for h in all_hits
                        if h.get("user", {}).get("name")).most_common(5)
    top_hosts = Counter(h.get("agent", {}).get("name","") for h in all_hits
                        if h.get("agent", {}).get("name")).most_common(5)
    countries = Counter(h.get("geo", {}).get("country","") for h in all_hits
                        if h.get("geo", {}).get("country")).most_common(5)
    return {
        "total_events": sum(q.get("total", 0) for q in queries_data),
        "total_hits":   len(all_hits),
        "flagged":      len(_flag_records(all_hits)),
        "severity":     sev_dist,
        "event_types":  et_dist,
        "top_ips":      top_ips,
        "top_users":    top_users,
        "top_hosts":    top_hosts,
        "countries":    countries,
    }


def _generate_overall_summary(queries_data: list[dict], stats: dict) -> str:
    try:
        from llm.client import ollama
        investigations = "\n".join(
            f"  {i}. '{q['natural_query']}' — {q.get('total',0):,} events, {q.get('time_range','—')}"
            for i, q in enumerate(queries_data, 1)
        )
        top3 = ", ".join(f"{k.replace('_',' ')} ({v})" for k, v in stats["event_types"].most_common(3))
        top_ips = ", ".join(ip for ip, _ in stats["top_ips"][:3]) or "none identified"
        prompt = f"""You are a senior security analyst writing the overall summary of a formal security report.
Write exactly 4 paragraphs — plain text, no bullets, no markdown, formal tone, third-person analyst voice.

Investigations covered:
{investigations}

Key statistics:
- Total events: {stats['total_events']:,}
- Flagged records: {stats['flagged']}
- Severity: {dict(stats['severity'])}
- Top attack types: {top3}
- Top source IPs: {top_ips}
- Overall risk: {_risk_rating(queries_data)}

Paragraphs:
1. Scope — what was investigated and the overall investigation coverage
2. Key findings — most significant threats, attack patterns, top offenders
3. Risk assessment — the organisation's current exposure based on this data
4. Response priority — high-level recommended immediate actions"""

        result = ollama.generate(prompt, temperature=0.2, max_tokens=450)
        if not result.startswith("ERROR:"):
            return result
    except Exception:
        pass

    # Structured fallback
    risk = _risk_rating(queries_data)
    top3 = [k.replace("_"," ") for k, _ in stats["event_types"].most_common(3)]
    return (
        f"This security report covers {len(queries_data)} investigation(s) conducted during "
        f"the current SIEM session, encompassing {stats['total_events']:,} security events "
        f"across the monitored environment.\n\n"
        f"The predominant threat categories were {', '.join(top3)}. "
        f"{stats['flagged']} records were automatically flagged for immediate attention "
        f"based on severity level, repeated source activity, or known attack signatures.\n\n"
        f"The overall risk posture is assessed as {risk}. "
        f"{stats['severity'].get('CRITICAL',0)} critical and "
        f"{stats['severity'].get('HIGH',0)} high-severity events were detected "
        f"in the sample reviewed.\n\n"
        f"Immediate review of flagged records is recommended. "
        f"Refer to the Recommended Actions section for prioritised remediation steps."
    )


def _build_report_charts(queries_data: list[dict], stats: dict) -> dict:
    """
    Build chart images for embedding in PDF/DOCX.
    Returns dict of {chart_name: PNG bytes}.

    Uses matplotlib — no Chrome/kaleido dependency.
    Dark theme to match the report's navy/accent palette.
    """
    import matplotlib
    matplotlib.use("Agg")  # non-interactive backend — safe for server use
    import matplotlib.pyplot as plt
    import matplotlib.patches as mpatches
    import numpy as np

    charts = {}

    # ── Colour palette matching the report ────────────────────────────────────
    BG      = "#0d1b2a"
    PANEL   = "#0f2035"
    TEXT    = "#c8d8f0"
    GRID    = "#1a3050"
    ACCENT  = "#0077cc"
    PALETTE = ["#0077cc","#00aaff","#e67e22","#c0392b","#27ae60","#8e44ad","#16a085"]
    SEV_C   = {"CRITICAL":"#c0392b","HIGH":"#e67e22","MEDIUM":"#0077cc","LOW":"#27ae60"}

    def _fig(w=7.5, h=4):
        fig, ax = plt.subplots(figsize=(w, h))
        fig.patch.set_facecolor(BG)
        ax.set_facecolor(PANEL)
        ax.tick_params(colors=TEXT, labelsize=8)
        ax.xaxis.label.set_color(TEXT)
        ax.yaxis.label.set_color(TEXT)
        for spine in ax.spines.values():
            spine.set_edgecolor(GRID)
        ax.grid(color=GRID, linewidth=0.5, alpha=0.7)
        return fig, ax

    def _save(fig):
        buf = io.BytesIO()
        fig.tight_layout()
        fig.savefig(buf, format="png", dpi=150, bbox_inches="tight",
                    facecolor=BG, edgecolor="none")
        plt.close(fig)
        buf.seek(0)
        return buf.read()

    # ── 1. Severity distribution (horizontal bar) ──────────────────────────────
    sev_order  = ["CRITICAL","HIGH","MEDIUM","LOW"]
    sev_counts = [stats["severity"].get(s,0) for s in sev_order]
    sev_colors = [SEV_C[s] for s in sev_order]

    if any(sev_counts):
        fig, ax = _fig(6, 3)
        bars = ax.barh(sev_order, sev_counts, color=sev_colors, height=0.5)
        ax.set_xlabel("Event Count", color=TEXT, fontsize=8)
        ax.set_title("Severity Distribution", color=TEXT, fontsize=10, pad=8)
        for bar, cnt in zip(bars, sev_counts):
            if cnt:
                ax.text(bar.get_width() + max(sev_counts)*0.01, bar.get_y() + bar.get_height()/2,
                        str(cnt), va="center", color=TEXT, fontsize=8)
        ax.set_xlim(0, max(sev_counts) * 1.15 if max(sev_counts) else 1)
        ax.invert_yaxis()
        charts["severity"] = _save(fig)

    # ── 2. Event type breakdown (pie / donut) ──────────────────────────────────
    et = stats["event_types"]
    if et:
        top_types  = et.most_common(7)
        labels     = [t.replace("_"," ").title() for t,_ in top_types]
        values     = [c for _,c in top_types]
        fig, ax    = _fig(6.5, 4.5)
        wedges, texts, autotexts = ax.pie(
            values, labels=None, autopct="%1.0f%%",
            colors=PALETTE[:len(values)],
            wedgeprops={"linewidth":2,"edgecolor":BG},
            pctdistance=0.78, startangle=90,
        )
        for at in autotexts:
            at.set_color(TEXT); at.set_fontsize(7.5)
        # Donut hole
        centre = plt.Circle((0,0), 0.52, fc=BG)
        ax.add_patch(centre)
        ax.set_title("Event Type Distribution", color=TEXT, fontsize=10, pad=8)
        ax.legend(wedges, labels, loc="lower center", ncol=2,
                  bbox_to_anchor=(0.5,-0.12), fontsize=7.5,
                  facecolor=PANEL, edgecolor=GRID,
                  labelcolor=TEXT)
        charts["event_types"] = _save(fig)

    # ── 3. Top source IPs (bar) ────────────────────────────────────────────────
    if stats["top_ips"]:
        ips    = [ip for ip,_ in stats["top_ips"]]
        counts = [c  for _,c  in stats["top_ips"]]
        fig, ax = _fig(7, 3.5)
        bars = ax.bar(range(len(ips)), counts,
                      color=[ACCENT]*len(ips), width=0.55)
        ax.set_xticks(range(len(ips)))
        ax.set_xticklabels(ips, rotation=15, ha="right", fontsize=7.5)
        ax.set_ylabel("Events", color=TEXT, fontsize=8)
        ax.set_title("Top Source IPs", color=TEXT, fontsize=10, pad=8)
        for bar, cnt in zip(bars, counts):
            ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.1,
                    str(cnt), ha="center", color=TEXT, fontsize=8)
        charts["top_ips"] = _save(fig)

    # ── 4. Per-investigation event counts (grouped bar) ────────────────────────
    if len(queries_data) > 1:
        names  = [f"Inv {i+1}" for i in range(len(queries_data))]
        totals = [q.get("total",0) for q in queries_data]
        fig, ax = _fig(max(5, len(queries_data)*1.5), 3.5)
        bars = ax.bar(names, totals, color=PALETTE[:len(names)], width=0.5)
        ax.set_ylabel("Total Events", color=TEXT, fontsize=8)
        ax.set_title("Events per Investigation", color=TEXT, fontsize=10, pad=8)
        for bar, cnt in zip(bars, totals):
            ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + max(totals)*0.01,
                    f"{cnt:,}", ha="center", color=TEXT, fontsize=8)
        charts["investigations"] = _save(fig)

    # ── 5. Top countries (horizontal bar) ─────────────────────────────────────
    if stats["countries"]:
        c_names  = [c for c,_ in stats["countries"]]
        c_counts = [n for _,n in stats["countries"]]
        fig, ax  = _fig(7, 3.5)
        bars = ax.barh(c_names, c_counts,
                       color=["#e67e22","#c0392b","#8e44ad","#16a085","#0077cc"][:len(c_names)],
                       height=0.5)
        ax.set_xlabel("Events", color=TEXT, fontsize=8)
        ax.set_title("Top Source Countries", color=TEXT, fontsize=10, pad=8)
        for bar, cnt in zip(bars, c_counts):
            ax.text(bar.get_width() + max(c_counts)*0.01,
                    bar.get_y() + bar.get_height()/2,
                    str(cnt), va="center", color=TEXT, fontsize=8)
        ax.invert_yaxis()
        charts["countries"] = _save(fig)

    return charts


# ── PDF ────────────────────────────────────────────────────────────────────────

def build_pdf_report(
    queries_data: list[dict],
    report_title: str  = "Security Incident Report",
    analyst_name: str  = "SIEM Intelligence Assistant",
    organisation: str  = "Organisation",
    output_path: str   = None,
) -> bytes:

    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        PageBreak, HRFlowable, KeepTogether, Image as RLImage,
    )
    from reportlab.platypus.flowables import Flowable
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY

    W, H = A4

    # Palette
    NAVY   = colors.HexColor("#0d1b2a")
    BLUE   = colors.HexColor("#1b3a6b")
    ACCENT = colors.HexColor("#0077cc")
    LIGHT  = colors.HexColor("#e8f0fe")
    RED    = colors.HexColor("#c0392b")
    ORANGE = colors.HexColor("#e67e22")
    GREEN  = colors.HexColor("#27ae60")
    GREY   = colors.HexColor("#6c757d")
    LGREY  = colors.HexColor("#f4f6f9")
    RULE   = colors.HexColor("#dde3ec")
    WHITE  = colors.white

    SEV = {"CRITICAL": RED, "HIGH": ORANGE, "MEDIUM": ACCENT, "LOW": GREEN}
    PRI = {"IMMEDIATE": RED, "SHORT-TERM": ORANGE, "LONG-TERM": ACCENT}

    generated_at = datetime.now().strftime("%d %B %Y, %H:%M UTC")
    risk         = _risk_rating(queries_data)
    risk_color   = SEV.get(risk, GREY)
    stats        = _compute_stats(queries_data)

    # Build charts — PNG bytes for embedding
    charts = {}
    try:
        charts = _build_report_charts(queries_data, stats)
    except Exception as e:
        pass  # charts are optional — report builds without them

    def embed_chart(name, width_cm=15, caption=None):
        """Return list of flowables to embed a named chart."""
        png = charts.get(name)
        if not png:
            return []
        items = [RLImage(io.BytesIO(png), width=width_cm*cm,
                         height=width_cm*cm*0.55)]
        if caption:
            items.append(Paragraph(caption, s_small))
        items.append(Spacer(1, 0.3*cm))
        return items

    buf = io.BytesIO()

    def on_page(canvas, doc):
        canvas.saveState()
        if doc.page > 1:
            canvas.setFillColor(NAVY)
            canvas.rect(0, H - 1.1*cm, W, 1.1*cm, fill=1, stroke=0)
            canvas.setFillColor(WHITE)
            canvas.setFont("Helvetica-Bold", 7.5)
            canvas.drawString(1.8*cm, H - 0.75*cm, report_title.upper())
            canvas.setFont("Helvetica", 7)
            canvas.drawRightString(W - 1.8*cm, H - 0.75*cm, "CONFIDENTIAL")
        # Footer
        canvas.setFillColor(LGREY)
        canvas.rect(0, 0, W, 0.8*cm, fill=1, stroke=0)
        canvas.setFillColor(GREY)
        canvas.setFont("Helvetica", 6.5)
        canvas.drawString(1.8*cm, 0.28*cm,
            f"SIEM Intelligence Report  ·  {organisation}  ·  {generated_at}")
        canvas.drawRightString(W - 1.8*cm, 0.28*cm, f"Page {doc.page}")
        canvas.restoreState()

    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        rightMargin=1.8*cm, leftMargin=1.8*cm,
        topMargin=1.6*cm,   bottomMargin=1.4*cm,
        title=report_title, author=analyst_name,
        onPage=on_page,
    )

    ss = getSampleStyleSheet()

    def sty(name, **kw):
        return ParagraphStyle(name, parent=ss["Normal"], **kw)

    s_cov_title = sty("ct", fontSize=30, leading=36, textColor=WHITE,
                       fontName="Helvetica-Bold")
    s_cov_sub   = sty("cs", fontSize=12, leading=16, textColor=colors.HexColor("#a0b8d8"))
    s_h1        = sty("h1", fontSize=13, leading=17, textColor=NAVY,
                       fontName="Helvetica-Bold", spaceBefore=12, spaceAfter=6)
    s_h2        = sty("h2", fontSize=10.5, leading=14, textColor=BLUE,
                       fontName="Helvetica-Bold", spaceBefore=8, spaceAfter=4)
    s_body      = sty("bd", fontSize=9, leading=13.5,
                       textColor=colors.HexColor("#1a1a2e"), spaceAfter=5,
                       alignment=TA_JUSTIFY)
    s_small     = sty("sm", fontSize=7.5, leading=11, textColor=GREY)
    s_code      = sty("cd", fontSize=6.5, leading=9.5, fontName="Courier",
                       textColor=colors.HexColor("#2d2d2d"), backColor=LGREY,
                       borderPad=4, leftIndent=4, rightIndent=4)
    s_toc       = sty("tc", fontSize=9.5, leading=14, textColor=BLUE)
    s_callout   = sty("cl", fontSize=9, leading=13.5, textColor=NAVY,
                       backColor=LIGHT, borderPad=8)

    story = []

    # ── COVER ──────────────────────────────────────────────────────────────────

    class Cover(Flowable):
        def draw(self):
            c = self.canv
            c.saveState()
            # Full navy background
            c.setFillColor(NAVY)
            c.rect(-1.8*cm, -H + 1.4*cm, W + 0.4*cm, H, fill=1, stroke=0)
            # Accent bar
            c.setFillColor(ACCENT)
            c.rect(-1.8*cm, H*0.36 - 1.4*cm, W + 0.4*cm, 0.3*cm, fill=1, stroke=0)
            c.restoreState()

    story.append(Cover())
    story.append(Spacer(1, 3.2*cm))

    # Classification banner
    cls_tbl = Table([["CONFIDENTIAL  —  SECURITY SENSITIVE DOCUMENT"]],
                    colWidths=[W - 3.6*cm])
    cls_tbl.setStyle(TableStyle([
        ("BACKGROUND",   (0,0),(-1,-1), RED),
        ("TEXTCOLOR",    (0,0),(-1,-1), WHITE),
        ("FONTNAME",     (0,0),(-1,-1), "Helvetica-Bold"),
        ("FONTSIZE",     (0,0),(-1,-1), 8.5),
        ("ALIGN",        (0,0),(-1,-1), "CENTER"),
        ("TOPPADDING",   (0,0),(-1,-1), 5),
        ("BOTTOMPADDING",(0,0),(-1,-1), 5),
    ]))
    story.append(cls_tbl)
    story.append(Spacer(1, 1.2*cm))
    story.append(Paragraph(report_title, s_cov_title))
    story.append(Spacer(1, 0.3*cm))
    story.append(Paragraph("Security Intelligence Report — SIEM Analysis", s_cov_sub))
    story.append(Spacer(1, 2.2*cm))

    # Metadata table
    mk = sty("mk", fontSize=7.5, textColor=colors.HexColor("#6090b8"),
              fontName="Helvetica-Bold")
    mv = sty("mv", fontSize=9,   textColor=WHITE)
    meta_rows = [
        ("DATE", generated_at), ("ANALYST", analyst_name),
        ("ORGANISATION", organisation),
        ("INVESTIGATIONS", str(len(queries_data))),
        ("TOTAL EVENTS", f"{stats['total_events']:,}"),
        ("RISK RATING", risk),
    ]
    mt = Table([[Paragraph(k, mk), Paragraph(v, mv)] for k, v in meta_rows],
               colWidths=[4.5*cm, 11*cm])
    mt.setStyle(TableStyle([
        ("TOPPADDING",    (0,0),(-1,-1), 4),
        ("BOTTOMPADDING", (0,0),(-1,-1), 4),
        ("LINEBELOW",     (0,0),(-1,-2), 0.3, colors.HexColor("#1e3050")),
    ]))
    mt.setStyle(TableStyle([
        ("TEXTCOLOR", (1,5),(1,5), risk_color),
        ("FONTNAME",  (1,5),(1,5), "Helvetica-Bold"),
    ]))
    story.append(mt)
    story.append(Spacer(1, 1.2*cm))

    # KPI boxes
    kpi_items = [
        (str(stats["total_events"]),                    "TOTAL EVENTS"),
        (str(stats["flagged"]),                         "FLAGGED"),
        (str(stats["severity"].get("CRITICAL", 0)),     "CRITICAL"),
        (str(stats["severity"].get("HIGH", 0)),         "HIGH"),
    ]
    kpi_cells = []
    for val, label in kpi_items:
        cell = Table(
            [[Paragraph(f'<font size="20" color="#00aaff"><b>{val}</b></font>', ss["Normal"])],
             [Paragraph(f'<font size="7" color="#708090">{label}</font>', ss["Normal"])]],
            colWidths=[3.5*cm]
        )
        cell.setStyle(TableStyle([
            ("ALIGN",         (0,0),(-1,-1), "CENTER"),
            ("BACKGROUND",    (0,0),(-1,-1), colors.HexColor("#0f1e30")),
            ("TOPPADDING",    (0,0),(-1,-1), 7),
            ("BOTTOMPADDING", (0,0),(-1,-1), 7),
        ]))
        kpi_cells.append(cell)
    kpi_row = Table([kpi_cells], colWidths=[3.5*cm]*4, hAlign="LEFT")
    kpi_row.setStyle(TableStyle([
        ("LEFTPADDING",  (0,0),(-1,-1), 3),
        ("RIGHTPADDING", (0,0),(-1,-1), 3),
    ]))
    story.append(kpi_row)
    story.append(PageBreak())

    # ── TABLE OF CONTENTS ─────────────────────────────────────────────────────

    story.append(Spacer(1, 0.5*cm))
    story.append(Paragraph("Contents", s_h1))
    story.append(HRFlowable(width="100%", thickness=2, color=ACCENT, spaceAfter=10))

    toc = [
        ("1", "Overall Summary"),
        ("2", "Executive Summary"),
        ("3", "Event Statistics"),
        ("4", "Event Data"),
        ("5", "Flagged Records"),
        ("6", "Recommended Actions"),
        ("A", "Appendix"),
    ]
    toc_tbl = Table(
        [[Paragraph(f"<b>{n}</b>", s_toc), Paragraph(t, s_toc)] for n, t in toc],
        colWidths=[1.2*cm, 15*cm]
    )
    toc_tbl.setStyle(TableStyle([
        ("TOPPADDING",    (0,0),(-1,-1), 5),
        ("BOTTOMPADDING", (0,0),(-1,-1), 5),
        ("LINEBELOW",     (0,0),(-1,-1), 0.3, RULE),
    ]))
    story.append(toc_tbl)
    story.append(PageBreak())

    # ── SECTION HEADER helper ─────────────────────────────────────────────────

    def sec_hdr(num, title, color=ACCENT):
        n_sty = sty(f"sn{num}", fontSize=13, leading=17, textColor=color,
                     fontName="Helvetica-Bold")
        t_sty = sty(f"st{num}", fontSize=13, leading=17, textColor=NAVY,
                     fontName="Helvetica-Bold")
        t = Table(
            [[Paragraph(num, n_sty), Paragraph(title, t_sty)]],
            colWidths=[1.2*cm, None]
        )
        t.setStyle(TableStyle([
            ("VALIGN",       (0,0),(-1,-1), "MIDDLE"),
            ("RIGHTPADDING", (0,0),(0,-1), 8),
            ("LINEBELOW",    (0,0),(-1,-1), 1.5, color),
            ("BOTTOMPADDING",(0,0),(-1,-1), 6),
            ("TOPPADDING",   (0,0),(-1,-1), 4),
        ]))
        return t

    def mini_tbl(rows, cw, hc=NAVY):
        t = Table(rows, colWidths=cw)
        t.setStyle(TableStyle([
            ("BACKGROUND",    (0,0),(-1,0), hc),
            ("TEXTCOLOR",     (0,0),(-1,0), WHITE),
            ("FONTNAME",      (0,0),(-1,0), "Helvetica-Bold"),
            ("FONTSIZE",      (0,0),(-1,-1), 8),
            ("TOPPADDING",    (0,0),(-1,-1), 3),
            ("BOTTOMPADDING", (0,0),(-1,-1), 3),
            ("ROWBACKGROUNDS",(0,1),(-1,-1), [LGREY, WHITE]),
            ("GRID",          (0,0),(-1,-1), 0.3, RULE),
        ]))
        return t

    # ── 1: OVERALL SUMMARY ────────────────────────────────────────────────────

    story.append(sec_hdr("1", "Overall Summary"))
    story.append(Spacer(1, 0.2*cm))

    summary_text = _generate_overall_summary(queries_data, stats)
    for para in summary_text.split("\n\n"):
        if para.strip():
            story.append(Paragraph(para.strip(), s_body))

    story.append(Spacer(1, 0.4*cm))

    # Risk callout
    sev_summary = "  ·  ".join(
        f"{s}: {stats['severity'].get(s,0)}" for s in ["CRITICAL","HIGH","MEDIUM","LOW"]
        if stats["severity"].get(s,0) > 0
    )
    callout = Table([[Paragraph(
        f'<b>Risk Rating: <font color="#{risk_color.hexval().lstrip("#")}">{risk}</font></b>'
        f'  —  {stats["total_events"]:,} events  ·  {stats["flagged"]} flagged  ·  '
        f'{sev_summary}',
        s_callout
    )]], colWidths=[W - 3.6*cm])
    callout.setStyle(TableStyle([
        ("BACKGROUND",    (0,0),(-1,-1), LIGHT),
        ("TOPPADDING",    (0,0),(-1,-1), 10),
        ("BOTTOMPADDING", (0,0),(-1,-1), 10),
        ("LEFTPADDING",   (0,0),(-1,-1), 14),
        ("LINEAFTER",     (0,0),(0,-1), 4, ACCENT),
    ]))
    story.append(callout)
    story.append(PageBreak())

    # ── 2: EXECUTIVE SUMMARY ──────────────────────────────────────────────────

    story.append(sec_hdr("2", "Executive Summary"))
    story.append(Spacer(1, 0.2*cm))

    for i, q in enumerate(queries_data, 1):
        story.append(KeepTogether([
            Paragraph(f"Investigation {i}: {q.get('natural_query','—')}", s_h2),
            Paragraph(
                f'<font color="#888888">Time range: {q.get("time_range","—")}  ·  '
                f'Events: {q.get("total",0):,}  ·  '
                f'Type: {q.get("event_type","—").replace("_"," ").title()}</font>',
                s_small
            ),
            Spacer(1, 4),
            Paragraph(
                q.get("narrative") or
                f"Found {q.get('total',0):,} {q.get('event_type','').replace('_',' ')} "
                f"events over {q.get('time_range','—')}.",
                s_body
            ),
        ]))
        story.append(HRFlowable(width="100%", thickness=0.3, color=RULE, spaceAfter=8))

    story.append(PageBreak())

    # ── 3: EVENT STATISTICS ───────────────────────────────────────────────────

    story.append(sec_hdr("3", "Event Statistics"))
    story.append(Spacer(1, 0.2*cm))

    # Severity table
    story.append(Paragraph("Severity Distribution", s_h2))
    ts = stats["total_hits"] or 1
    sev_rows = [["Severity", "Count", "% of Sample"]]
    for sev in ["CRITICAL","HIGH","MEDIUM","LOW"]:
        cnt = stats["severity"].get(sev, 0)
        sev_rows.append([sev, str(cnt), f"{cnt/ts*100:.1f}%"])
    sev_tbl = mini_tbl(sev_rows, [5*cm, 4*cm, 4*cm])
    for si, sev in enumerate(["CRITICAL","HIGH","MEDIUM","LOW"], 1):
        sev_tbl.setStyle(TableStyle([
            ("TEXTCOLOR", (0,si),(0,si), SEV.get(sev,GREY)),
            ("FONTNAME",  (0,si),(0,si), "Helvetica-Bold"),
        ]))
    story.append(sev_tbl)
    story.append(Spacer(1, 0.4*cm))

    # Two-column offenders
    left, right = [], []
    if stats["top_ips"]:
        left += [Paragraph("Top Source IPs", s_h2),
                 mini_tbl([["IP Address","Events"]] + [[ip,str(c)] for ip,c in stats["top_ips"]],
                           [6.5*cm, 2.5*cm])]
    if stats["countries"]:
        left += [Spacer(1,0.4*cm),
                 Paragraph("Top Source Countries", s_h2),
                 mini_tbl([["Country","Events"]] + [[c,str(n)] for c,n in stats["countries"]],
                           [6.5*cm, 2.5*cm])]
    if stats["top_users"]:
        right += [Paragraph("Top Affected Users", s_h2),
                  mini_tbl([["User","Events"]] + [[u,str(c)] for u,c in stats["top_users"]],
                            [6.5*cm, 2.5*cm])]
    if stats["top_hosts"]:
        right += [Spacer(1,0.4*cm),
                  Paragraph("Top Affected Hosts", s_h2),
                  mini_tbl([["Host","Events"]] + [[h,str(c)] for h,c in stats["top_hosts"]],
                            [6.5*cm, 2.5*cm])]
    if left or right:
        pad = max(len(left), len(right))
        left  += [Spacer(1,1)] * (pad - len(left))
        right += [Spacer(1,1)] * (pad - len(right))
        cols = Table([[left, right]], colWidths=[9.2*cm, 9.2*cm])
        cols.setStyle(TableStyle([
            ("VALIGN", (0,0),(-1,-1), "TOP"),
            ("LEFTPADDING", (1,0),(1,-1), 10),
        ]))
        story.append(cols)

    # Embedded charts
    story.append(Spacer(1, 0.4*cm))
    if charts:
        story.append(Paragraph("Visual Summaries", s_h2))
        story.append(Spacer(1, 0.2*cm))

        # Severity and event type side by side
        sev_png = charts.get("severity")
        et_png  = charts.get("event_types")
        if sev_png and et_png:
            chart_row = Table([[
                RLImage(io.BytesIO(sev_png), width=8.8*cm, height=4.8*cm),
                RLImage(io.BytesIO(et_png),  width=8.8*cm, height=4.8*cm),
            ]], colWidths=[9.1*cm, 9.1*cm])
            chart_row.setStyle(TableStyle([
                ("LEFTPADDING",   (1,0),(1,-1), 6),
                ("RIGHTPADDING",  (0,0),(0,-1), 6),
                ("VALIGN",        (0,0),(-1,-1), "TOP"),
            ]))
            story.append(chart_row)
            story.append(Spacer(1, 0.2*cm))

        # Top IPs and countries side by side
        ip_png  = charts.get("top_ips")
        c_png   = charts.get("countries")
        if ip_png and c_png:
            chart_row2 = Table([[
                RLImage(io.BytesIO(ip_png), width=8.8*cm, height=4.2*cm),
                RLImage(io.BytesIO(c_png),  width=8.8*cm, height=4.2*cm),
            ]], colWidths=[9.1*cm, 9.1*cm])
            chart_row2.setStyle(TableStyle([
                ("LEFTPADDING",   (1,0),(1,-1), 6),
                ("RIGHTPADDING",  (0,0),(0,-1), 6),
                ("VALIGN",        (0,0),(-1,-1), "TOP"),
            ]))
            story.append(chart_row2)
            story.append(Spacer(1, 0.2*cm))
        elif ip_png:
            story += embed_chart("top_ips", width_cm=14)
        elif c_png:
            story += embed_chart("countries", width_cm=14)

        # Investigations comparison (only if > 1)
        inv_png = charts.get("investigations")
        if inv_png:
            story.append(Spacer(1, 0.2*cm))
            story += embed_chart("investigations", width_cm=12,
                                  caption="Events per investigation across this session.")

    story.append(PageBreak())

    # ── 4: EVENT DATA TABLES ──────────────────────────────────────────────────

    story.append(sec_hdr("4", "Event Data Tables"))
    story.append(Spacer(1, 0.2*cm))

    for i, q in enumerate(queries_data, 1):
        hits = q.get("hits", [])
        if not hits:
            continue
        story.append(Paragraph(f"Investigation {i}: {q.get('natural_query','—')}", s_h2))
        story.append(Paragraph(
            f'<font color="#888888">Showing {min(len(hits),30)} of {q.get("total",0):,} '
            f'events  ·  {q.get("time_range","—")}</font>', s_small))
        story.append(Spacer(1, 4))

        hdrs = ["Timestamp", "Event Type", "Description", "User", "Src IP", "Sev"]
        rows = [hdrs] + [
            [h.get("@timestamp","")[:16],
             h.get("event_type","—"),
             h.get("rule",{}).get("description","—")[:42],
             h.get("user",{}).get("name","—") or "—",
             h.get("data",{}).get("srcip","—") or "—",
             _severity_label(h.get("rule",{}).get("level",""))]
            for h in hits[:30]
        ]
        cw  = [3.4*cm, 3*cm, 5.2*cm, 2.5*cm, 3*cm, 1.7*cm]
        tbl = Table(rows, colWidths=cw, repeatRows=1)
        ts2 = [
            ("BACKGROUND",    (0,0),(-1,0), BLUE),
            ("TEXTCOLOR",     (0,0),(-1,0), WHITE),
            ("FONTNAME",      (0,0),(-1,0), "Helvetica-Bold"),
            ("FONTSIZE",      (0,0),(-1,-1), 7.5),
            ("ALIGN",         (0,0),(-1,-1), "LEFT"),
            ("TOPPADDING",    (0,0),(-1,-1), 3),
            ("BOTTOMPADDING", (0,0),(-1,-1), 3),
            ("ROWBACKGROUNDS",(0,1),(-1,-1), [WHITE, LGREY]),
            ("GRID",          (0,0),(-1,-1), 0.3, RULE),
        ]
        for ri, row in enumerate(rows[1:], 1):
            sc = SEV.get(row[5], GREY)
            ts2 += [("TEXTCOLOR",(5,ri),(5,ri),sc), ("FONTNAME",(5,ri),(5,ri),"Helvetica-Bold")]
        tbl.setStyle(TableStyle(ts2))
        story.append(tbl)
        if q.get("total",0) > 30:
            story.append(Paragraph(
                f"Showing first 30 of {q.get('total',0):,} events. "
                "Full dataset available in SIEM.", s_small))
        story.append(Spacer(1, 0.5*cm))

    story.append(PageBreak())

    # ── 5: FLAGGED RECORDS ────────────────────────────────────────────────────

    story.append(sec_hdr("5", "Flagged Records", color=RED))
    story.append(Spacer(1, 0.2*cm))

    all_hits = [h for q in queries_data for h in q.get("hits", [])]
    flagged  = _flag_records(all_hits)

    if flagged:
        story.append(Paragraph(
            f'{len(flagged)} record(s) require immediate analyst attention — '
            f'flagged by severity level, repeated source activity, or attack signatures.',
            s_body))
        story.append(Spacer(1, 0.2*cm))
        fhdrs = ["Timestamp","Type","Description","Src IP","Sev","Flag Reason"]
        frows = [fhdrs] + [
            [h.get("@timestamp","")[:16],
             h.get("event_type","—"),
             h.get("rule",{}).get("description","—")[:36],
             h.get("data",{}).get("srcip","—") or "—",
             _severity_label(h.get("rule",{}).get("level","")),
             ", ".join(h.get("_flag_reasons",[]))[:40]]
            for h in flagged[:25]
        ]
        cw  = [3.2*cm, 2.8*cm, 4.2*cm, 2.5*cm, 1.8*cm, 4.3*cm]
        ft  = Table(frows, colWidths=cw, repeatRows=1)
        fts = [
            ("BACKGROUND",    (0,0),(-1,0), colors.HexColor("#8b0000")),
            ("TEXTCOLOR",     (0,0),(-1,0), WHITE),
            ("FONTNAME",      (0,0),(-1,0), "Helvetica-Bold"),
            ("FONTSIZE",      (0,0),(-1,-1), 7.5),
            ("TOPPADDING",    (0,0),(-1,-1), 3),
            ("BOTTOMPADDING", (0,0),(-1,-1), 3),
            ("ROWBACKGROUNDS",(0,1),(-1,-1), [colors.HexColor("#fff5f5"), WHITE]),
            ("GRID",          (0,0),(-1,-1), 0.3, colors.HexColor("#ffcccc")),
            ("TEXTCOLOR",     (5,1),(5,-1), RED),
        ]
        for ri, row in enumerate(frows[1:], 1):
            sc = SEV.get(row[4], GREY)
            fts += [("TEXTCOLOR",(4,ri),(4,ri),sc),("FONTNAME",(4,ri),(4,ri),"Helvetica-Bold")]
        ft.setStyle(TableStyle(fts))
        story.append(ft)
    else:
        story.append(Paragraph(
            "No records met the automatic flagging criteria. "
            "This does not indicate absence of risk — manual review is always recommended.",
            s_body))

    story.append(PageBreak())

    # ── 6: RECOMMENDED ACTIONS ────────────────────────────────────────────────

    story.append(sec_hdr("6", "Recommended Actions", color=ORANGE))
    story.append(Spacer(1, 0.2*cm))
    story.append(Paragraph(
        "Actions are categorised as IMMEDIATE (within 24 hours), "
        "SHORT-TERM (within 1–2 weeks), or LONG-TERM (within 90 days).",
        s_body))
    story.append(Spacer(1, 0.3*cm))

    seen = list(dict.fromkeys(q.get("event_type","default") for q in queries_data))
    for et in seen:
        recs = _RECS.get(et, _RECS["default"])
        story.append(Paragraph(et.replace("_"," ").title(), s_h2))
        rec_rows = []
        for priority, action in recs:
            pc = PRI.get(priority, GREY)
            rec_rows.append([
                Paragraph(f'<font color="#{pc.hexval().lstrip("#")}"><b>{priority}</b></font>',
                           sty(f"rp{priority}", fontSize=7.5, fontName="Helvetica-Bold")),
                Paragraph(action, sty("ra", fontSize=8.5, leading=12)),
            ])
        rt = Table(rec_rows, colWidths=[3*cm, 14*cm])
        rt.setStyle(TableStyle([
            ("TOPPADDING",    (0,0),(-1,-1), 4),
            ("BOTTOMPADDING", (0,0),(-1,-1), 4),
            ("LINEBELOW",     (0,0),(-1,-2), 0.3, RULE),
            ("VALIGN",        (0,0),(-1,-1), "MIDDLE"),
        ]))
        story.append(rt)
        story.append(Spacer(1, 0.2*cm))

    story.append(PageBreak())

    # ── APPENDIX ──────────────────────────────────────────────────────────────

    story.append(sec_hdr("A", "Appendix — Queries Executed", color=GREY))
    story.append(Spacer(1, 0.2*cm))
    story.append(Paragraph(
        "The Elasticsearch DSL queries below were used to retrieve the data in this report.",
        s_body))
    story.append(Spacer(1, 0.3*cm))

    for i, q in enumerate(queries_data, 1):
        story.append(Paragraph(f"Query {i}: {q.get('natural_query','—')}", s_h2))
        dsl = q.get("dsl")
        if dsl:
            dsl_str = json.dumps(dsl.get("body",{}), indent=2)
            if len(dsl_str) > 1000:
                dsl_str = dsl_str[:1000] + "\n... (truncated)"
            story.append(Paragraph(
                dsl_str.replace("&","&amp;").replace("<","&lt;")
                       .replace("\n","<br/>").replace(" ","&nbsp;"),
                s_code))
        story.append(Spacer(1, 0.3*cm))

    doc.build(story)
    pdf_bytes = buf.getvalue()
    if output_path:
        with open(output_path, "wb") as f:
            f.write(pdf_bytes)
    return pdf_bytes


# ── DOCX ───────────────────────────────────────────────────────────────────────

def build_docx_report(
    queries_data: list[dict],
    report_title: str = "Security Incident Report",
    analyst_name: str = "SIEM Intelligence Assistant",
    organisation: str = "Organisation",
    output_path: str  = None,
) -> bytes:
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
    except ImportError:
        raise ImportError("python-docx not installed. Run: pip install python-docx")

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Cm(2.5)
        sec.left_margin = sec.right_margin = Cm(2.5)

    generated_at = datetime.now().strftime("%d %B %Y, %H:%M UTC")
    risk         = _risk_rating(queries_data)
    stats        = _compute_stats(queries_data)
    RC = {"CRITICAL":RGBColor(192,0,0),"HIGH":RGBColor(230,126,34),
          "MEDIUM":RGBColor(0,119,204),"LOW":RGBColor(39,174,96)}
    PR = {"IMMEDIATE":RGBColor(192,0,0),"SHORT-TERM":RGBColor(230,126,34),
          "LONG-TERM":RGBColor(0,119,204)}
    risk_rgb = RC.get(risk, RGBColor(100,100,100))

    # Build charts
    charts = {}
    try:
        charts = _build_report_charts(queries_data, stats)
    except Exception:
        pass

    def h(text, lvl=1, color=None, bold=True):
        hd = doc.add_heading(text, level=lvl)
        for r in hd.runs:
            if color: r.font.color.rgb = color
            r.bold = bold

    def p(text, bold=False, sz=10, color=None, italic=False):
        para = doc.add_paragraph()
        r    = para.add_run(text)
        r.bold = bold; r.italic = italic
        r.font.size = Pt(sz)
        if color: r.font.color.rgb = color
        return para

    # Cover
    doc.add_heading(report_title, 0)
    p("Security Intelligence Report — SIEM Analysis", sz=12,
      color=RGBColor(100,100,120), italic=True)
    doc.add_paragraph()
    meta = doc.add_table(rows=6, cols=2)
    meta.style = "Table Grid"
    for i,(k,v) in enumerate([
        ("Date",generated_at),("Analyst",analyst_name),("Organisation",organisation),
        ("Investigations",str(len(queries_data))),
        ("Total Events",f"{stats['total_events']:,}"),("Risk Rating",risk),
    ]):
        meta.rows[i].cells[0].text = k
        meta.rows[i].cells[1].text = v
        meta.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        if k == "Risk Rating":
            meta.rows[i].cells[1].paragraphs[0].runs[0].font.color.rgb = risk_rgb
            meta.rows[i].cells[1].paragraphs[0].runs[0].bold = True
    doc.add_page_break()

    # Overall Summary
    h("1  Overall Summary", 1, color=RGBColor(13,27,42))
    overall = _generate_overall_summary(queries_data, stats)
    for para in overall.split("\n\n"):
        if para.strip(): p(para.strip())
    sev_line = "  ·  ".join(
        f"{s}: {stats['severity'].get(s,0)}"
        for s in ["CRITICAL","HIGH","MEDIUM","LOW"] if stats["severity"].get(s,0)
    )
    p(f"Overall Risk: {risk}  —  {stats['total_events']:,} events  ·  "
      f"{stats['flagged']} flagged  ·  {sev_line}",
      bold=True, color=risk_rgb)
    doc.add_page_break()

    # Executive Summary
    h("2  Executive Summary", 1, color=RGBColor(13,27,42))
    for i,q in enumerate(queries_data,1):
        h(f"Investigation {i}: {q.get('natural_query','—')}", 2)
        p(f"Time: {q.get('time_range','—')}  ·  Events: {q.get('total',0):,}  ·  "
          f"Type: {q.get('event_type','—').replace('_',' ').title()}",
          sz=8, color=RGBColor(108,117,125), italic=True)
        p(q.get("narrative") or f"Found {q.get('total',0):,} events.")
    doc.add_page_break()

    # Statistics
    h("3  Event Statistics", 1, color=RGBColor(13,27,42))
    h("Severity Distribution", 2)
    t = doc.add_table(rows=1, cols=3); t.style="Table Grid"
    for c,lbl in zip(t.rows[0].cells,["Severity","Count","% of Sample"]):
        c.text=lbl; c.paragraphs[0].runs[0].bold=True
        c.paragraphs[0].runs[0].font.size=Pt(8)
    total_s = stats["total_hits"] or 1
    for sev in ["CRITICAL","HIGH","MEDIUM","LOW"]:
        cnt = stats["severity"].get(sev,0)
        row = t.add_row().cells
        row[0].text=sev; row[1].text=str(cnt); row[2].text=f"{cnt/total_s*100:.1f}%"
        row[0].paragraphs[0].runs[0].font.color.rgb = RC.get(sev,RGBColor(100,100,100))
        row[0].paragraphs[0].runs[0].bold=True
        for c in row:
            for para in c.paragraphs:
                for r in para.runs: r.font.size=Pt(8)
    if stats["top_ips"]:
        h("Top Source IPs",2)
        t2=doc.add_table(rows=1,cols=2); t2.style="Table Grid"
        t2.rows[0].cells[0].text="IP Address"; t2.rows[0].cells[1].text="Events"
        for ip,cnt in stats["top_ips"]:
            r=t2.add_row().cells; r[0].text=ip; r[1].text=str(cnt)

    # Embed charts into docx
    if charts:
        doc.add_paragraph()
        h("Visual Summaries", 2)
        for chart_name, caption in [
            ("severity",       "Severity Distribution"),
            ("event_types",    "Event Type Breakdown"),
            ("top_ips",        "Top Source IPs"),
            ("countries",      "Top Source Countries"),
            ("investigations", "Events per Investigation"),
        ]:
            png = charts.get(chart_name)
            if png:
                try:
                    para = doc.add_paragraph()
                    para.add_run(caption).bold = True
                    doc.add_picture(io.BytesIO(png), width=Cm(14))
                    doc.add_paragraph()
                except Exception:
                    pass

    doc.add_page_break()

    # Event Data
    h("4  Event Data Tables", 1, color=RGBColor(13,27,42))
    for i,q in enumerate(queries_data,1):
        hits=q.get("hits",[]); 
        if not hits: continue
        h(f"Investigation {i}: {q.get('natural_query','—')}",2)
        p(f"Showing {min(len(hits),25)} of {q.get('total',0):,}  ·  {q.get('time_range','—')}",
          sz=8,italic=True,color=RGBColor(108,117,125))
        tbl=doc.add_table(rows=1,cols=6); tbl.style="Table Grid"
        for c,lbl in zip(tbl.rows[0].cells,["Time","Type","Description","User","Src IP","Sev"]):
            c.text=lbl; c.paragraphs[0].runs[0].bold=True
            c.paragraphs[0].runs[0].font.size=Pt(8)
        for hit in hits[:25]:
            row=tbl.add_row().cells
            row[0].text=hit.get("@timestamp","")[:16]
            row[1].text=hit.get("event_type","—")
            row[2].text=hit.get("rule",{}).get("description","—")[:40]
            row[3].text=hit.get("user",{}).get("name","—") or "—"
            row[4].text=hit.get("data",{}).get("srcip","—") or "—"
            sev=_severity_label(hit.get("rule",{}).get("level",""))
            row[5].text=sev
            row[5].paragraphs[0].runs[0].font.color.rgb=RC.get(sev,RGBColor(100,100,100))
            row[5].paragraphs[0].runs[0].bold=True
            for c in row:
                for para in c.paragraphs:
                    for r in para.runs: r.font.size=Pt(8)
        doc.add_paragraph()
    doc.add_page_break()

    # Flagged Records
    h("5  Flagged Records", 1, color=RGBColor(192,0,0))
    all_hits=[h2 for q in queries_data for h2 in q.get("hits",[])]
    flagged=_flag_records(all_hits)
    if flagged:
        p(f"{len(flagged)} record(s) flagged.",bold=True,color=RGBColor(192,0,0))
        for hit in flagged[:20]:
            reasons=", ".join(hit.get("_flag_reasons",[]))
            p(f"• {hit.get('@timestamp','')[:16]} | {hit.get('event_type','—')} | "
              f"{hit.get('rule',{}).get('description','—')[:38]} | "
              f"IP: {hit.get('data',{}).get('srcip','—')} | {reasons}",
              sz=9, color=RGBColor(192,0,0))
    else:
        p("No records met the flagging criteria.")
    doc.add_page_break()

    # Recommended Actions
    h("6  Recommended Actions", 1, color=RGBColor(13,27,42))
    p("IMMEDIATE = 24h  ·  SHORT-TERM = 2 weeks  ·  LONG-TERM = 90 days",
      sz=8,italic=True,color=RGBColor(108,117,125))
    doc.add_paragraph()
    for et in dict.fromkeys(q.get("event_type","default") for q in queries_data):
        h(et.replace("_"," ").title(), 2)
        for priority, action in _RECS.get(et,_RECS["default"]):
            para = doc.add_paragraph()
            r1=para.add_run(f"[{priority}]  "); r1.bold=True; r1.font.size=Pt(9)
            r1.font.color.rgb=PR.get(priority,RGBColor(100,100,100))
            r2=para.add_run(action); r2.font.size=Pt(9)
    doc.add_page_break()

    # Appendix
    h("A  Appendix — Queries Executed", 1, color=RGBColor(100,100,100))
    for i,q in enumerate(queries_data,1):
        h(f"Query {i}: {q.get('natural_query','—')}",2)
        if q.get("dsl"):
            para=doc.add_paragraph()
            r=para.add_run(json.dumps(q["dsl"].get("body",{}),indent=2)[:600])
            r.font.name="Courier New"; r.font.size=Pt(7.5)

    buf=io.BytesIO(); doc.save(buf); b=buf.getvalue()
    if output_path:
        with open(output_path,"wb") as f: f.write(b)
    return b


# ── Convenience ────────────────────────────────────────────────────────────────

def build_report_from_session(session_log, connector, selected_indices=None,
                               fmt="pdf", report_title=None) -> bytes:
    if not session_log: raise ValueError("No session log entries.")
    records = session_log
    if selected_indices:
        records = [r for r in session_log if r.index in selected_indices]
    if not records: raise ValueError("No matching entries.")
    queries_data = []
    for rec in records:
        if not rec.dsl: continue
        result = connector.execute(rec.dsl, allow_widen=True)
        queries_data.append({
            "natural_query": rec.natural_query,
            "hits":          result.get("hits",[]),
            "aggregations":  result.get("aggregations",{}),
            "narrative":     None,
            "time_range":    rec.time_range,
            "total":         result.get("total",0),
            "event_type":    rec.intent,
            "dsl":           rec.dsl,
        })
    if not queries_data: raise ValueError("No data fetched.")
    title = report_title or f"Security Report — {datetime.now().strftime('%Y-%m-%d')}"
    if fmt == "pdf":  return build_pdf_report(queries_data, report_title=title)
    if fmt == "docx": return build_docx_report(queries_data, report_title=title)
    raise ValueError(f"Unknown format: {fmt}")